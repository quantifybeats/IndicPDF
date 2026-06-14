import re
import statistics
import unicodedata
from collections import Counter
from pathlib import Path
from docx import Document
from docx.shared import Pt
from pdfminer.high_level import extract_pages
from pdfminer.layout import LAParams, LTTextBoxHorizontal, LTTextLineHorizontal, LTChar

try:
    from .font_manager import font_registry
    from .encoding_manager import encoding_manager
except (ImportError, ValueError):
    from font_manager import font_registry
    from encoding_manager import encoding_manager

def normalize_font_name(font_name: str) -> str:
    if not font_name:
        return "Normal"
    match = re.search(r'[A-Z]{6}\+(.+)', font_name)
    return match.group(1) if match else font_name


# Preferred installed family per script, so the emitted DOCX names a font the
# reader actually has (a PDF subset name like 'ABCDEF+Foo' opens as tofu in Word).
_SCRIPT_DOCX_FONT = {
    "devanagari": "Noto Sans Devanagari",
    "bengali":    "Noto Sans Bengali",
    "gurmukhi":   "Noto Sans Gurmukhi",
    "gujarati":   "Noto Sans Gujarati",
    "odia":       "Noto Sans Oriya",
    "tamil":      "Noto Sans Tamil",
    "telugu":     "Noto Sans Telugu",
    "kannada":    "Noto Sans Kannada",
    "malayalam":  "Noto Sans Malayalam",
}

# (registry_script, unicode_start, unicode_end) for the Indic blocks we handle.
_SCRIPT_RANGES = [
    ("devanagari", 0x0900, 0x097F), ("bengali", 0x0980, 0x09FF),
    ("gurmukhi",   0x0A00, 0x0A7F), ("gujarati", 0x0A80, 0x0AFF),
    ("odia",       0x0B00, 0x0B7F), ("tamil",    0x0B80, 0x0BFF),
    ("telugu",     0x0C00, 0x0C7F), ("kannada",  0x0C80, 0x0CFF),
    ("malayalam",  0x0D00, 0x0D7F),
]
# Danda/double-danda are shared Indic punctuation in the Devanagari block — they
# must not decide the script (see processor._detect_script).
_SHARED_INDIC_PUNCT = {0x0964, 0x0965}

# De-fragmentation tuning. A horizontal gap wider than WORD_GAP * font-size is a
# real word space; narrower gaps between positioned glyphs are spurious and are
# dropped. A vertical gap wider than PARA_GAP * line-height starts a new paragraph.
_WORD_GAP = 0.28
_PARA_GAP = 1.4
_LINE_TOL = 0.6   # fraction of glyph height within which glyphs share a line


# Broken-ToUnicode garbage signature. A corrupt ToUnicode CMap maps every glyph
# to the same (or a tiny set of) wrong code point, so extraction yields a wall of
# one repeated akshara ('సససస', 'తెతెతె', 'नननन'). Those are valid Indic code
# points — they survive NFC and junk-stripping — so extraction cannot recover the
# real text; only OCR or fixing the source can. Detect and refuse rather than
# emit a confidently-wrong DOCX.
_GARBAGE_MIN_LEN = 24       # too short to judge below this
_GARBAGE_MAX_DISTINCT = 4   # real Indic text this long uses far more aksharas
_GARBAGE_DOMINANCE = 0.40   # one char covering this share = repeated-glyph signature


class ToUnicodeGarbageError(RuntimeError):
    """The source PDF's text layer is corrupt (broken ToUnicode): glyphs all map
    to the same wrong character, so no extractor can recover the real text."""


def _indic_chars(text: str):
    """Indic letters/signs in text, excluding the shared danda punctuation.
    The garbage check scores on these alone so Latin in a mixed-language doc
    can't dilute a broken Indic run below the threshold (real case:
    'English ... नननन')."""
    out = []
    for ch in text:
        cp = ord(ch)
        if cp in _SHARED_INDIC_PUNCT:
            continue
        for _reg, lo, hi in _SCRIPT_RANGES:
            if lo <= cp <= hi:
                out.append(ch)
                break
    return out


def _is_tounicode_garbage(text: str) -> bool:
    """True when the Indic text shows the broken-ToUnicode signature: a long run
    of aksharas collapsed onto one (or a few) repeated code points. Scored on the
    Indic subset only, so Latin and normal Indic text return False — and a broken
    Indic run survives dilution by surrounding Latin."""
    indic = _indic_chars(text)
    if len(indic) < _GARBAGE_MIN_LEN:
        return False
    if len(set(indic)) <= _GARBAGE_MAX_DISTINCT:
        return True
    top = Counter(indic).most_common(1)[0][1]
    return (top / len(indic)) >= _GARBAGE_DOMINANCE


def _detect_doc_script(text: str):
    """Dominant Indic script in text (ignoring the shared danda), or None."""
    counts = {}
    for ch in text:
        cp = ord(ch)
        if cp in _SHARED_INDIC_PUNCT:
            continue
        for reg, lo, hi in _SCRIPT_RANGES:
            if lo <= cp <= hi:
                counts[reg] = counts.get(reg, 0) + 1
                break
    return max(counts, key=counts.get) if counts else None


def _set_run_fonts(run, name):
    """Set the font on a run for ALL script classes. python-docx's run.font.name
    only writes w:ascii/w:hAnsi, but Indic text is a *complex script* — Word and
    LibreOffice pick the w:cs font for it. Without w:cs the reader renders Telugu
    etc. with the default complex-script font (detached matras / wrong shaping).
    Set w:cs (and w:eastAsia) too so the chosen font is used everywhere."""
    from docx.oxml.ns import qn
    run.font.name = name  # w:ascii + w:hAnsi
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), name)


def _assign_docx_font(run, clean_text, source_font, script):
    """Name a font the reader has: an installed script font for Indic runs,
    otherwise the resolved family for the run's original font. The name is
    applied to the complex-script slot too, so Indic shaping survives in Word."""
    pref = _SCRIPT_DOCX_FONT.get(script) if script else None
    if pref and font_registry.get_font_metadata(pref):
        _set_run_fonts(run, pref)
        return
    metadata = font_registry.get_font_metadata(source_font)
    if metadata:
        _set_run_fonts(run, metadata.family_name)
    elif clean_text.strip():
        fallback = font_registry.resolve_font(
            source_font, ord(clean_text.strip()[0]), script=script)
        if fallback:
            _set_run_fonts(run, fallback.stem)


def _iter_glyphs(obj):
    """Recurse a pdfminer layout object, yielding non-blank LTChar glyphs."""
    if isinstance(obj, LTChar):
        if obj.get_text() and not obj.get_text().isspace():
            yield obj
    elif hasattr(obj, "__iter__"):
        for child in obj:
            yield from _iter_glyphs(child)


def _page_is_fragmented(page_layout) -> bool:
    """True when the page contains Indic text whose extracted form is shattered
    into tiny whitespace-separated tokens — the signature of a glyph-positioned
    PDF where every akshara was placed individually (pdfminer then inserts a
    space between each). Catches both manifestations: one cluster per box, and
    many clusters per box separated by spurious spaces. Latin/normal PDFs return
    False and keep the fast line-level path."""
    indic = "".join(
        b.get_text() for b in page_layout
        if isinstance(b, LTTextBoxHorizontal) and _detect_doc_script(b.get_text())
    )
    if not indic.strip():
        return False
    tokens = [t for t in re.split(r'\s+', indic) if t]
    if len(tokens) < 8:
        return False
    tiny = sum(1 for t in tokens if len(t) <= 2)
    return (tiny / len(tokens)) >= 0.5


def _reflow_page(page_layout):
    """Rebuild logical paragraphs from raw glyph positions, bypassing pdfminer's
    broken box/line grouping. Returns a list of paragraphs, each a list of line
    dicts {text, size, font}. Every glyph is preserved; only spurious inter-
    akshara spaces and per-glyph line breaks are removed."""
    glyphs = list(_iter_glyphs(page_layout))
    if not glyphs:
        return []
    heights = [g.y1 - g.y0 for g in glyphs if g.y1 > g.y0]
    med_h = statistics.median(heights) if heights else 10.0

    # Cluster glyphs into visual lines by y-midpoint, top to bottom.
    glyphs.sort(key=lambda g: (-((g.y0 + g.y1) / 2), g.x0))
    line_groups, cur, cur_mid = [], [], None
    tol = _LINE_TOL * med_h
    for g in glyphs:
        mid = (g.y0 + g.y1) / 2
        if cur_mid is None or abs(mid - cur_mid) <= tol:
            cur.append(g)
            cur_mid = sum((x.y0 + x.y1) / 2 for x in cur) / len(cur)
        else:
            line_groups.append(cur)
            cur, cur_mid = [g], mid
    if cur:
        line_groups.append(cur)

    # Build each line: join glyphs, inserting a space only at real word gaps.
    # The threshold is adaptive — a word gap must beat both a fixed fraction of
    # the font size AND a multiple of this line's typical (intra-word) glyph gap,
    # so uniformly wide-spaced aksharas don't each become a "word".
    lines = []
    for grp in line_groups:
        grp.sort(key=lambda g: g.x0)
        size = statistics.median([g.size for g in grp]) or 12.0
        gaps = [b.x0 - a.x1 for a, b in zip(grp, grp[1:]) if b.x0 - a.x1 > 0]
        med_gap = statistics.median(gaps) if gaps else 0.0
        word_threshold = max(_WORD_GAP * size, med_gap * 1.8)
        parts, prev = [], None
        for g in grp:
            if prev is not None and (g.x0 - prev.x1) > word_threshold:
                parts.append(" ")
            parts.append(g.get_text())
            prev = g
        font = normalize_font_name(
            next((g.fontname for g in grp if getattr(g, "fontname", None)), "Normal"))
        lines.append({
            "text": "".join(parts), "size": round(size, 1), "font": font,
            "ytop": max(g.y1 for g in grp), "ybot": min(g.y0 for g in grp),
        })

    # Group lines into paragraphs by vertical gap.
    lhs = [ln["ytop"] - ln["ybot"] for ln in lines if ln["ytop"] > ln["ybot"]]
    med_lh = statistics.median(lhs) if lhs else med_h
    paras, cur = [], [lines[0]]
    for prev, ln in zip(lines, lines[1:]):
        if (prev["ybot"] - ln["ytop"]) > _PARA_GAP * med_lh:
            paras.append(cur)
            cur = [ln]
        else:
            cur.append(ln)
    paras.append(cur)
    return paras

def _emit_linewise_page(doc, page_layout):
    """Existing path: trust pdfminer's box/line grouping (correct for normal,
    non-glyph-positioned PDFs such as Latin documents)."""
    boxes = sorted([b for b in page_layout if isinstance(b, LTTextBoxHorizontal)],
                   key=lambda b: b.y1, reverse=True)
    for box in boxes:
        paragraph = doc.add_paragraph()
        for line in box:
            if not isinstance(line, LTTextLineHorizontal):
                continue
            line_text = line.get_text()

            # Pre-extract font name for CID resolution context.
            font_name = "Normal"
            for char in line:
                if hasattr(char, 'fontname'):
                    font_name = normalize_font_name(char.fontname)
                    break

            def cid_replacer(match, _fn=font_name):
                cid_val = int(match.group(1))
                resolved = encoding_manager.resolve_cid(cid_val, _fn)
                return resolved if resolved is not None else match.group(0)

            processed_line = re.sub(r'\(cid\s*:\s*(\d+)\s*\)', cid_replacer, line_text)
            processed_line = re.sub(r'cid\s*:\s*(\d+)', cid_replacer, processed_line)

            clean_text = encoding_manager.strip_all_junk(processed_line)
            clean_text = unicodedata.normalize('NFC', clean_text)
            if not clean_text.strip():
                continue

            script = _detect_doc_script(clean_text)

            font_size = 12
            for char in line:
                if hasattr(char, 'size'):
                    font_size = round(char.size, 1)
                    break

            run = paragraph.add_run(clean_text)
            _assign_docx_font(run, clean_text, font_name, script)
            run.font.size = Pt(font_size)


def _emit_reflowed_page(doc, page_layout):
    """De-fragmented path: rebuild logical paragraphs from raw glyph positions so
    a glyph-positioned Indic PDF no longer becomes one akshara per paragraph with
    spurious inter-akshara spaces. Characters are preserved verbatim."""
    for para_lines in _reflow_page(page_layout):
        text = " ".join(ln["text"] for ln in para_lines)
        clean = unicodedata.normalize('NFC', encoding_manager.strip_all_junk(text))
        if not clean.strip():
            continue
        script = _detect_doc_script(clean)
        size = round(statistics.median([ln["size"] for ln in para_lines]), 1)
        source_font = para_lines[0]["font"]
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(clean)
        _assign_docx_font(run, clean, source_font, script)
        run.font.size = Pt(size)


def process_pdf_to_docx(pdf_path: Path, docx_output_path: Path):
    """PDF to DOCX extractor. Normal PDFs use pdfminer's line grouping; glyph-
    positioned Indic PDFs (which pdfminer shatters into one cluster per box) are
    re-flowed from raw glyph coordinates so the output keeps readable spacing and
    paragraphs without resorting to lossy OCR."""
    report = {"pages_processed": 0, "reflowed_pages": 0, "status": "success"}
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(72); section.bottom_margin = Pt(72)
        section.left_margin = Pt(72); section.right_margin = Pt(72)

    try:
        laparams = LAParams(word_margin=0.1, line_margin=0.5, char_margin=2.0)
        pages = list(extract_pages(pdf_path, laparams=laparams))
        report["pages_processed"] = len(pages)

        for page_layout in pages:
            if _page_is_fragmented(page_layout):
                report["reflowed_pages"] += 1
                _emit_reflowed_page(doc, page_layout)
            else:
                _emit_linewise_page(doc, page_layout)

        # Refuse to ship a confidently-wrong DOCX: if extraction collapsed the
        # whole document onto one repeated akshara, the source ToUnicode is
        # corrupt and only OCR (or fixing the source) can recover it.
        if _is_tounicode_garbage("\n".join(p.text for p in doc.paragraphs)):
            raise ToUnicodeGarbageError(
                "Source PDF text layer is corrupt (broken ToUnicode: every glyph "
                "maps to the same character, e.g. 'సససస'). Text extraction cannot "
                "recover it. Use OCR mode, or regenerate the source PDF."
            )

        doc.save(docx_output_path)
    except Exception as e:
        report["status"] = "error"
        raise e

    return report
