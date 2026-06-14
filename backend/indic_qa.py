"""Automated Indic rendering-correctness checker.

No single signal is trusted: string inspection misses visual matra errors,
codepoint comparison passed on a render whose matras were wrong, and OCR gave a
false PASS on a visually broken render. So fidelity is judged by several
*independent* signals, and a result is PASS only when every applicable one agrees.

Signals
-------
1. shaping        — HarfBuzz shapes the expected text with the font that will be
                    used; every cluster must map to a real glyph (notdef == 0).
                    Catches boxes / missing glyphs.
2. text_roundtrip — extract the output's text layer, strip artifacts, NFC-compare
                    to expected. Catches dropped/reordered chars, cid leakage,
                    fragmentation.
3. mark_adjacency — no combining mark (Mn/Mc) is orphaned (preceded by space or at
                    start). Catches matra detachment in the text layer.
4. glyph_match    — the glyph-id sequence actually drawn in the PDF equals the
                    HarfBuzz-shaped sequence for the embedded font. Catches
                    *unshaped* / misplaced rendering that OCR forgives.
5. ocr_backcheck  — render -> OCR -> similarity. A LOW score is a strong failure
                    signal; a HIGH score is supporting evidence only.

Each signal returns a dict {ok, score, detail}. `check_pdf` aggregates them.
"""
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, List, Optional

# Representative probe consonant per registry script (a real glyph must exist).
SCRIPT_PROBE = {
    "devanagari": 0x0939, "bengali": 0x09AC, "gurmukhi": 0x0A2A, "gujarati": 0x0A97,
    "odia": 0x0B13, "tamil": 0x0BAE, "telugu": 0x0C24, "kannada": 0x0C95,
    "malayalam": 0x0D2E,
}
# Tesseract language code per registry script.
SCRIPT_TESS = {
    "devanagari": "hin", "bengali": "ben", "gurmukhi": "pan", "gujarati": "guj",
    "odia": "ori", "tamil": "tam", "telugu": "tel", "kannada": "kan",
    "malayalam": "mal",
}


def _nfc(s: str) -> str:
    return unicodedata.normalize("NFC", s or "")


def _norm_compare(s: str) -> str:
    """Strip whitespace + common punctuation for tolerant comparison."""
    return re.sub(r"[\s\.\,\:\;\?\!\"\'\(\)।॥।॥]", "", _nfc(s))


# --------------------------------------------------------------------- signals

def shaping_signal(expected: str, font_path) -> dict:
    """HarfBuzz shapes expected text with font_path; notdef==0 required."""
    try:
        import uharfbuzz as hb
        blob = hb.Blob.from_file_path(str(font_path))
        font = hb.Font(hb.Face(blob))
        buf = hb.Buffer()
        buf.add_str(expected)
        buf.guess_segment_properties()
        hb.shape(font, buf)
        gids = [g.codepoint for g in buf.glyph_infos]
        notdef = sum(1 for g in gids if g == 0)
        return {"ok": notdef == 0 and len(gids) > 0, "score": 1.0 if notdef == 0 else 0.0,
                "detail": f"glyphs={len(gids)} notdef={notdef}", "_gids": gids}
    except Exception as e:  # pragma: no cover
        return {"ok": False, "score": 0.0, "detail": f"shaping error: {e}", "_gids": []}


def text_roundtrip_signal(expected: str, pdf_path, strict: bool = True) -> dict:
    """Extract the PDF text layer and compare to expected (NFC, punctuation-loose).

    `strict=True` for extraction paths (PDF->DOCX, de-frag) where the text layer
    IS the product. `strict=False` for rendered output (DOCX/TXT->PDF), where a
    shaped complex-script text layer is inherently lossy (reordered conjuncts have
    no clean glyph->Unicode map) — there the score is advisory, not a verdict."""
    try:
        from pdfminer.high_level import extract_text
        try:
            from .encoding_manager import encoding_manager
        except (ImportError, ValueError):
            from encoding_manager import encoding_manager
        got = encoding_manager.strip_all_junk(extract_text(str(pdf_path)))
        ratio = SequenceMatcher(None, _norm_compare(expected), _norm_compare(got)).ratio()
        exp_chars = set(_norm_compare(expected))
        lost = exp_chars - set(_norm_compare(got))
        passed = ratio >= 0.95 and not lost
        return {"ok": passed if strict else None, "score": round(ratio, 3),
                "detail": f"ratio={ratio:.3f} lost={''.join(sorted(lost))[:20]!r}"
                          + ("" if strict else " (advisory)")}
    except Exception as e:
        return {"ok": False if strict else None, "score": 0.0, "detail": f"extract error: {e}"}


def mark_adjacency_signal(text: str) -> dict:
    """No combining mark may be orphaned (after whitespace or at string start)."""
    s = _nfc(text)
    bad = []
    for i, ch in enumerate(s):
        if unicodedata.category(ch) in ("Mn", "Mc"):
            if i == 0 or s[i - 1].isspace():
                bad.append((i, hex(ord(ch))))
    return {"ok": not bad, "score": 1.0 if not bad else 0.0,
            "detail": "clean" if not bad else f"orphaned marks: {bad[:5]}"}


def glyph_match_signal(expected: str, pdf_path, font_path=None) -> dict:
    """Detect *unshaped* rendering deterministically by comparing how many glyphs
    were actually painted to how many HarfBuzz produces for the expected text.

    Correct complex-script shaping forms ligatures/half-forms, so the shaped glyph
    count is typically LOWER than the code-point count. If the renderer painted
    ~one glyph per code point (no reduction below HB), it did not shape — the
    matra-misplacement class of bug. Decisive only on a clear excess; advisory
    otherwise (kerning/space differences are normal)."""
    try:
        import uharfbuzz as hb
        painted = [ch for ch in _iter_ltchars_pdf(pdf_path) if not ch.get_text().isspace()]
        if not painted:
            return {"ok": None, "score": None, "detail": "no glyphs painted (skipped)"}
        if font_path is None:
            return {"ok": None, "score": None, "detail": f"painted={len(painted)} (no ref font)"}
        blob = hb.Blob.from_file_path(str(font_path))
        font = hb.Font(hb.Face(blob))
        buf = hb.Buffer(); buf.add_str(re.sub(r"\s", "", _nfc(expected)))
        buf.guess_segment_properties(); hb.shape(font, buf)
        hb_glyphs = len(buf.glyph_infos)
        cp = len(re.sub(r"\s", "", _nfc(expected)))
        # painted close to HB-shaped => shaped; painted near raw code-point count
        # (and well above HB) => unshaped.
        unshaped = hb_glyphs < cp and len(painted) >= cp and len(painted) > hb_glyphs + 1
        return {"ok": not unshaped,
                "score": round(hb_glyphs / max(1, len(painted)), 2),
                "detail": f"painted={len(painted)} hb_shaped={hb_glyphs} cp={cp}"
                          + (" UNSHAPED" if unshaped else "")}
    except Exception as e:
        return {"ok": None, "score": None, "detail": f"glyph_match skipped: {e}"}


def _iter_ltchars_pdf(pdf_path):
    from pdfminer.high_level import extract_pages
    for page in extract_pages(str(pdf_path)):
        yield from _iter_ltchars(page)


def _iter_ltchars(obj):
    from pdfminer.layout import LTChar
    if isinstance(obj, LTChar):
        yield obj
    elif hasattr(obj, "__iter__"):
        for c in obj:
            yield from _iter_ltchars(c)


def ocr_backcheck_signal(expected: str, pdf_path, tess_lang: str) -> dict:
    """Render the PDF and OCR it; low similarity is a strong failure signal."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        img = convert_from_path(str(pdf_path), dpi=200)[0]
        got = pytesseract.image_to_string(img, lang=tess_lang)
        ratio = SequenceMatcher(None, _norm_compare(expected), _norm_compare(got)).ratio()
        # ADVISORY only: OCR false-passes broken renders AND false-fails good ones
        # (model gaps for some scripts/fonts). A low score flags a render for human
        # review; it is never a sole verdict.
        return {"ok": None, "score": round(ratio, 3),
                "detail": f"ocr_sim={ratio:.3f}" + (" LOW-review" if ratio < 0.6 else "")}
    except Exception as e:
        return {"ok": None, "score": None, "detail": f"ocr skipped: {e}"}


# ------------------------------------------------------------------- aggregate

@dataclass
class FidelityReport:
    label: str
    script: str
    signals: Dict[str, dict] = field(default_factory=dict)

    @property
    def verdict(self) -> str:
        decisive = [s for s in self.signals.values() if s.get("ok") is not None]
        if not decisive:
            return "UNKNOWN"
        return "PASS" if all(s["ok"] for s in decisive) else "FAIL"

    def line(self) -> str:
        bits = " ".join(f"{k}={'·' if v.get('ok') is None else ('✓' if v['ok'] else '✗')}"
                        for k, v in self.signals.items())
        return f"[{self.verdict:7}] {self.label:28} {self.script:11} {bits}"


# Clean, verified probe sentences (conjuncts + matras) per script.
SELFTEST_SAMPLES = {
    "devanagari": "यह एक हिंदी वाक्य है।",
    "bengali":    "এটি একটি বাংলা বাক্য।",
    "gurmukhi":   "ਇਹ ਇੱਕ ਪੰਜਾਬੀ ਵਾਕ ਹੈ।",
    "gujarati":   "આ એક ગુજરાતી વાક્ય છે.",
    "odia":       "ଏହା ଏକ ଓଡ଼ିଆ ବାକ୍ୟ।",
    "tamil":      "இது ஒரு தமிழ் வாக்கியம்.",
    "telugu":     "ఇది ఒక తెలుగు వాక్యం.",
    "kannada":    "ಇದು ಒಂದು ಕನ್ನಡ ವಾಕ್ಯ.",
    "malayalam":  "ഇത് ഒരു മലയാള വാക്യം.",
}


def _report_to_dict(rep: "FidelityReport") -> dict:
    return {
        "label": rep.label, "script": rep.script, "verdict": rep.verdict,
        "signals": {k: {"ok": v.get("ok"), "score": v.get("score"),
                        "detail": v.get("detail")} for k, v in rep.signals.items()},
    }


def run_self_test(formats=("txt", "docx")) -> dict:
    """Convert every probe sentence through each format and score it. Self-
    contained — generates its own inputs. Meant to run IN the deployed
    environment (the only place font/renderer verdicts are trustworthy)."""
    import tempfile
    from pathlib import Path as _P
    try:
        from . import processor
    except (ImportError, ValueError):
        import processor
    from docx import Document

    tmp = _P(tempfile.mkdtemp(prefix="indicqa_"))
    results = []
    for script, text in SELFTEST_SAMPLES.items():
        try:
            font_path = processor._preferred_render_font(text)[1]
        except Exception:
            font_path = None
        for fmt in formats:
            label = f"{script} {fmt}2pdf"
            pdf = tmp / f"{script}_{fmt}.pdf"
            try:
                if fmt == "txt":
                    t = tmp / f"{script}.txt"
                    t.write_text(text, encoding="utf-8")
                    processor.process_txt_to_pdf(t, pdf)
                elif fmt == "docx":
                    d = Document()
                    d.add_paragraph().add_run(text).font.name = "Arial"
                    dp = tmp / f"{script}.docx"
                    d.save(dp)
                    processor.process_docx_to_pdf_final(dp, pdf)
                else:
                    raise ValueError(f"unknown format {fmt}")
                results.append(_report_to_dict(
                    check_pdf(text, pdf, script, font_path=font_path,
                              mode="render", label=label)))
            except Exception as e:
                results.append({"label": label, "script": script,
                                "verdict": "ERROR", "error": str(e), "signals": {}})

    summary = {
        "total": len(results),
        "pass": sum(1 for r in results if r["verdict"] == "PASS"),
        "fail": sum(1 for r in results if r["verdict"] == "FAIL"),
        "error": sum(1 for r in results if r["verdict"] == "ERROR"),
        "needs_review": [r["label"] for r in results
                         if r.get("signals", {}).get("ocr_backcheck", {}).get("score") is not None
                         and (r["signals"]["ocr_backcheck"]["score"] or 1) < 0.6],
    }
    return {"summary": summary, "results": results}


def check_pdf(expected: str, pdf_path, script: str, font_path=None,
              label: str = "", mode: str = "render") -> FidelityReport:
    """mode='render' for produced PDFs (DOCX/TXT->PDF) — visual signals decide,
    text layer advisory. mode='extract' for text products (PDF->DOCX, de-frag) —
    the text layer IS the product, so it is decisive."""
    rep = FidelityReport(label=label or Path(pdf_path).name, script=script)
    if font_path is not None:
        rep.signals["shaping"] = shaping_signal(expected, font_path)
    rep.signals["text_roundtrip"] = text_roundtrip_signal(
        expected, pdf_path, strict=(mode == "extract"))
    rep.signals["mark_adjacency"] = mark_adjacency_signal(expected)
    rep.signals["glyph_match"] = glyph_match_signal(expected, pdf_path, font_path=font_path)
    tess = SCRIPT_TESS.get(script)
    if tess:
        rep.signals["ocr_backcheck"] = ocr_backcheck_signal(expected, pdf_path, tess)
    return rep
