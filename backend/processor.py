import logging
import unicodedata
import subprocess
import zipfile
import shutil
import os
import xml.etree.ElementTree as ET
from pathlib import Path
from pypdf import PdfReader
from fpdf import FPDF
from docx import Document

class IndicPDF(FPDF):
    def header(self):
        pass
    def footer(self):
        pass

logger = logging.getLogger(__name__)

try:
    from .font_manager import font_registry
    from .encoding_manager import encoding_manager
except (ImportError, ValueError):
    from font_manager import font_registry
    from encoding_manager import encoding_manager

def extract_docx_fonts(docx_path: Path):
    """Extract required font names from DOCX fontTable.xml."""
    fonts = set()
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/fontTable.xml' in z.namelist():
                xml_content = z.read('word/fontTable.xml')
                root = ET.fromstring(xml_content)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for font in root.findall('.//w:font', ns):
                    name = font.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                    if name:
                        fonts.add(name)
    except Exception as e:
        logger.warning(f"Failed to extract fonts from DOCX: {e}")
    return fonts

_INDIC_KEYWORDS = {
    "devanagari", "hindi", "bengali", "gujarati", "telugu", "tamil",
    "kannada", "malayalam", "punjabi", "odia", "oriya", "sanskrit",
    "mangal", "kokila", "utsaah", "aparajita", "lohit", "noto sans dev",
    "noto serif dev", "noto sans ben", "noto sans guj", "noto sans tel",
    "noto sans tam", "noto sans kan", "noto sans mal",
}

def _is_indic_font(font_name: str) -> bool:
    n = font_name.lower()
    return any(kw in n for kw in _INDIC_KEYWORDS)


# --- Indic fidelity helpers (defense-in-depth, shared by every render path) ---
# (fpdf_script_tag, registry_script_name, unicode_start, unicode_end)
_SCRIPT_RANGES = [
    ("deva", "devanagari", 0x0900, 0x097F),
    ("beng", "bengali",    0x0980, 0x09FF),
    ("guru", "gurmukhi",   0x0A00, 0x0A7F),
    ("gujr", "gujarati",   0x0A80, 0x0AFF),
    ("orya", "odia",       0x0B00, 0x0B7F),
    ("taml", "tamil",      0x0B80, 0x0BFF),
    ("telu", "telugu",     0x0C00, 0x0C7F),
    ("knda", "kannada",    0x0C80, 0x0CFF),
    ("mlym", "malayalam",  0x0D00, 0x0D7F),
]

# Preferred installed family per registry script. Forced onto Indic-bearing runs
# so a Latin font name on Indic text cannot trigger a tofu substitution. NOTE:
# the NotoSans{Tamil,Bengali,Gujarati,Kannada,Malayalam,Oriya} files in the repo
# are broken ~21KB placeholders with no real glyphs, so the Lohit family (verified
# to cover each script) is the curated choice for those. Devanagari/Telugu have
# real Noto fonts. The pick is coverage-checked below regardless of this table.
_SCRIPT_RENDER_FONT = {
    "devanagari": "Noto Sans Devanagari",
    "telugu":     "Noto Sans Telugu",
    "tamil":      "Lohit Tamil",
    "bengali":    "Lohit Bengali",
    "gujarati":   "Lohit Gujarati",
    "kannada":    "Lohit Kannada",
    "malayalam":  "Lohit Malayalam",
    "odia":       "Lohit Odia",
}

# One representative consonant per script, used to verify a candidate font really
# contains the script's glyphs (guards against broken placeholder fonts).
_SCRIPT_PROBE_CP = {
    "devanagari": 0x0939, "telugu": 0x0C24, "tamil": 0x0BAE, "bengali": 0x09AC,
    "gujarati": 0x0A97, "kannada": 0x0C95, "malayalam": 0x0D2E, "odia": 0x0B13,
}

# reg_script -> chosen FontMetadata (or None); resolved once per process.
_render_font_cache = {}


def _font_covers(path, codepoint: int) -> bool:
    """True if the font's cmap actually maps codepoint (not just a name match)."""
    try:
        from fontTools.ttLib import TTFont
        return codepoint in TTFont(str(path)).getBestCmap()
    except Exception:
        return False


# Danda (U+0964) and double danda (U+0965) live in the Devanagari block but are
# shared punctuation across Indic scripts — they must NOT decide the script.
_SHARED_INDIC_PUNCT = {0x0964, 0x0965}


def _detect_script(text: str):
    """Return (fpdf_tag, registry_script) for the *dominant* Indic script in the
    text, or (None, None). Picks the script with the most letters so a shared
    danda can't misclassify e.g. Bengali or Odia as Devanagari."""
    counts = {}
    for ch in text:
        cp = ord(ch)
        if cp in _SHARED_INDIC_PUNCT:
            continue
        for fpdf_tag, reg_name, lo, hi in _SCRIPT_RANGES:
            if lo <= cp <= hi:
                counts[(fpdf_tag, reg_name)] = counts.get((fpdf_tag, reg_name), 0) + 1
                break
    if not counts:
        return None, None
    return max(counts.items(), key=lambda kv: kv[1])[0]


def _clean_render_text(text: str) -> str:
    """Strip PDF extraction artifacts and NFC-normalize text bound for a renderer."""
    return unicodedata.normalize("NFC", encoding_manager.strip_all_junk(text or ""))


def _preferred_render_font_meta(reg):
    """Return a FontMetadata whose font actually covers `reg`'s glyphs, or None.

    Tries the curated family first, then every font mapped to the script,
    verifying real glyph coverage so broken placeholder fonts are skipped."""
    if not reg:
        return None
    if reg in _render_font_cache:
        return _render_font_cache[reg]
    probe = _SCRIPT_PROBE_CP.get(reg)
    candidates = []
    pref = _SCRIPT_RENDER_FONT.get(reg)
    if pref:
        m = font_registry.get_font_metadata(pref)
        if m:
            candidates.append(m)
    for m in font_registry.script_fallback.get(reg, []):
        if m not in candidates:
            candidates.append(m)
    chosen = None
    for m in candidates:
        if probe is None or _font_covers(m.path, probe):
            chosen = m
            break
    _render_font_cache[reg] = chosen
    return chosen


def _preferred_render_font(text: str):
    """(registry_script, Path) of a coverage-verified font for the text's script,
    else (registry_script or None, None)."""
    _, reg = _detect_script(text)
    if not reg:
        return None, None
    meta = _preferred_render_font_meta(reg)
    return reg, (meta.path if meta else None)


def _iter_doc_paragraphs(doc):
    """Yield every paragraph in document order, including those inside tables."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for block in doc.element.body:
        if block.tag == qn("w:p"):
            yield Paragraph(block, doc)
        elif block.tag == qn("w:tbl"):
            for row in Table(block, doc).rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _sanitize_docx_for_render(src: Path) -> Path:
    """Return a DOCX safe to hand to LibreOffice headless: extraction junk
    stripped from every run, and Indic-bearing runs forced onto an installed
    script font so they cannot be substituted into tofu (boxes).

    Returns the original path untouched when nothing needed changing.
    """
    doc = Document(src)
    changed = False
    for para in _iter_doc_paragraphs(doc):
        for run in para.runs:
            cleaned = _clean_render_text(run.text)
            if cleaned != run.text:
                run.text = cleaned
                changed = True
            if not cleaned.strip():
                continue
            _, reg_script = _detect_script(cleaned)
            meta = _preferred_render_font_meta(reg_script)
            if meta:
                desired = meta.family_name
                if run.font.name != desired:
                    run.font.name = desired
                    changed = True
    if not changed:
        return src
    out = src.with_name(f"{src.stem}_sanitized.docx")
    doc.save(out)
    return out

def check_fonts_availability(requested_fonts):
    """Check only Indic fonts — LibreOffice handles standard Western fonts natively."""
    missing = []
    for font_name in requested_fonts:
        if not font_name or any(x in font_name.lower() for x in ["minor", "major", "theme"]):
            continue
        if not _is_indic_font(font_name):
            continue
        if not font_registry.get_font_metadata(font_name):
            missing.append(font_name)
    return missing

def setup_libreoffice_fonts():
    """Ensure all registry fonts are visible to fontconfig/LibreOffice."""
    fonts_dir = Path.home() / ".fonts"
    fonts_dir.mkdir(exist_ok=True)
    
    updated = False
    for name, metadata in font_registry.registry.items():
        dest = fonts_dir / metadata.path.name
        if not dest.exists():
            try:
                os.symlink(metadata.path, dest)
                updated = True
            except:
                try:
                    shutil.copy(metadata.path, dest)
                    updated = True
                except: pass
    
    if updated:
        try:
            subprocess.run(["fc-cache", "-f", "-v"], capture_output=True, check=False)
        except: pass

def extract_pdf_fonts(pdf_path: Path):
    """Extract font names from a generated PDF for validation."""
    fonts = set()
    try:
        reader = PdfReader(str(pdf_path))
        for page in reader.pages:
            if '/Resources' in page and '/Font' in page['/Resources']:
                font_resources = page['/Resources']['/Font']
                for font_key in font_resources:
                    font_obj = font_resources[font_key].get_object()
                    base_font = font_obj.get('/BaseFont', '')
                    if base_font:
                        # Strip subset prefix (e.g., 'ABCDEF+Arial' -> 'Arial')
                        name = str(base_font).split('+')[-1] if '+' in str(base_font) else str(base_font)
                        fonts.add(name.replace('/', ''))
    except Exception as e:
        logger.warning(f"Failed to extract PDF fonts: {e}")
    return fonts

# QA F12: the FPDF fallback loads the entire DOCX XML via python-docx and
# renders run-by-run with HarfBuzz shaping. A small (~90 KB) upload whose
# word/document.xml decompresses to tens of MB (a zip-amplification bomb that
# passes the 25 MB *compressed* upload cap) then pins a CPU for minutes and,
# with one worker, starves the whole queue. Bound it the same way the
# reconstruction engine does (QA F5).
MAX_DOCX_XML_BYTES = 20 * 1024 * 1024   # decompressed word/document.xml
MAX_RENDER_PARAGRAPHS = 5000


def _assert_docx_within_render_limits(docx_path: Path) -> None:
    """Reject decompression bombs before python-docx loads the whole tree."""
    try:
        with zipfile.ZipFile(docx_path) as archive:
            info = archive.getinfo("word/document.xml")
    except KeyError:
        return  # no document.xml — let Document() raise its own clear error
    if info.file_size > MAX_DOCX_XML_BYTES:
        raise ValueError(
            f"DOCX content too large to render: word/document.xml decompresses "
            f"to {info.file_size // (1024 * 1024)} MB (limit "
            f"{MAX_DOCX_XML_BYTES // (1024 * 1024)} MB). Split the document and retry."
        )


def process_docx_to_pdf_fpdf_fallback(docx_path: Path, pdf_output_path: Path):
    """Fallback docx to pdf conversion using fpdf2 directly."""
    logger.warning("LibreOffice headless not found. Falling back to native FPDF rendering.")
    report = {
        "font_substitutions": [],
        "legacy_conversions": [],
        "legacy_unsupported": [],
        "warnings": [],
        "status": "success",
        "engine": "FPDF-Native-Fallback"
    }
    _assert_docx_within_render_limits(docx_path)
    doc = Document(docx_path)
    pdf = IndicPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    registered_fonts = set()

    def _render_paragraph(para):
        for run in para.runs:
            if run.text != " " and not run.text.strip():
                continue

            requested_font = run.font.name or "Normal"

            # Strip PDF text-extraction artifacts (cid markers, U+FFFD boxes)
            # before anything renders them verbatim.
            processed_text = _clean_render_text(run.text)

            encoding_type = encoding_manager.detect_legacy_encoding(requested_font)
            if encoding_type:
                if encoding_manager.legacy_supported(encoding_type):
                    processed_text = encoding_manager.convert_to_unicode(processed_text, encoding_type)
                    report["legacy_conversions"].append({"font": requested_font, "type": encoding_type})
                else:
                    # QA F6: a legacy 8-bit Indic font we can't faithfully map.
                    # Don't claim a conversion — the output for this run is
                    # unreliable; surface it instead of silently garbling.
                    entry = {"font": requested_font, "type": encoding_type}
                    if entry not in report["legacy_unsupported"]:
                        report["legacy_unsupported"].append(entry)
                        report["warnings"].append(
                            f"Unsupported legacy encoding '{encoding_type}' "
                            f"(font '{requested_font}'); text rendered as-is and may be garbled."
                        )

            # Detect script up front so font resolution is script-aware. The
            # prebuilt font index carries no per-glyph unicode_ranges, so the
            # char-only fallback in resolve_font can't match — without an
            # explicit script, Indic runs resolve to nothing and crash on
            # helvetica (QA F3).
            script, registry_script = _detect_script(processed_text)

            # For Indic runs, pick a coverage-verified script font first (skips
            # broken placeholder fonts); otherwise honour the requested font.
            resolved_path = None
            if registry_script:
                _, resolved_path = _preferred_render_font(processed_text)
            if not resolved_path:
                resolved_path = font_registry.resolve_font(requested_font, script=registry_script)
            if not resolved_path and processed_text:
                resolved_path = font_registry.resolve_font(
                    requested_font, ord(processed_text[0]), script=registry_script
                )

            if resolved_path:
                res_name = resolved_path.stem
                req_norm = requested_font.lower().replace(" ", "").replace("-", "")
                res_norm = res_name.lower().replace(" ", "").replace("-", "")
                if req_norm not in res_norm and res_norm not in req_norm and requested_font != "Normal":
                    report["font_substitutions"].append({"requested": requested_font, "resolved": res_name})

                font_id = resolved_path.stem
                if font_id not in registered_fonts:
                    try:
                        pdf.add_font(font_id, "", str(resolved_path))
                        registered_fonts.add(font_id)
                    except Exception as e:
                        logger.error(f"Failed to add font {font_id}: {e}")

                pdf.set_font(font_id, size=12)
                pdf.set_text_shaping(use_shaping_engine=bool(script), script=script, direction="ltr")
                pdf.write(h=10, text=processed_text)
            else:
                # No font covers this run. helvetica only handles latin-1, so
                # guard the write — a single unsupported script must not abort
                # the whole document (QA F3/F10).
                pdf.set_font("helvetica", size=12)
                try:
                    pdf.write(h=10, text=processed_text)
                except Exception as e:
                    logger.warning(f"Skipped unrenderable run ({requested_font}): {e}")
                    report.setdefault("unrendered_runs", []).append({
                        "font": requested_font, "reason": "no_font_covers_script",
                    })

        pdf.ln(10)

    # Render paragraphs and tables in document order
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    rendered = 0
    truncated = False
    for block in doc.element.body:
        if rendered >= MAX_RENDER_PARAGRAPHS:
            truncated = True
            break
        if block.tag == qn("w:p"):
            _render_paragraph(Paragraph(block, doc))
            rendered += 1
        elif block.tag == qn("w:tbl"):
            tbl = Table(block, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _render_paragraph(para)
                        rendered += 1

    if truncated:
        # Don't silently drop content (QA F12): cap and flag it.
        report["warnings"].append(
            f"Document exceeded {MAX_RENDER_PARAGRAPHS} paragraphs; output truncated."
        )
        report["truncated"] = True

    pdf.output(str(pdf_output_path))
    return report

def process_docx_to_pdf_final(docx_path: Path, pdf_output_path: Path):
    """
    LibreOffice Headless Pipeline with FPDF fallback.
    Strictly preserves fonts, layout, and styling.
    """
    # Check if soffice is available
    soffice_path = "soffice"
    if shutil.which("soffice") is None:
        mac_soffice = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        if mac_soffice.exists():
            soffice_path = str(mac_soffice)
        else:
            soffice_path = None

    if soffice_path is None:
        return process_docx_to_pdf_fpdf_fallback(docx_path, pdf_output_path)

    report = {"status": "success", "engine": "LibreOffice-Headless"}

    # 0. Sanitize for render: strip PDF extraction artifacts (cid markers,
    # U+FFFD) from every run and force Indic-bearing runs onto an installed
    # script font so LibreOffice cannot substitute them into tofu (boxes).
    # Returns the original path when nothing needed fixing.
    render_src = _sanitize_docx_for_render(docx_path)

    # 1. Font Metadata Extraction (from the sanitized doc — its runs now request
    # the fonts we will actually render with).
    docx_fonts = extract_docx_fonts(render_src)
    logger.info(f"DOCX requested fonts: {docx_fonts}")
    
    # 2. Strict Font Validation
    missing = check_fonts_availability(docx_fonts)
    if missing:
        error_msg = f"Strict font preservation failed. Missing fonts: {', '.join(missing)}. Please upload these fonts to the server registry to ensure 100% fidelity."
        logger.error(error_msg)
        raise ValueError(error_msg)

    # 3. Synchronize Registry with System Fontconfig
    setup_libreoffice_fonts()

    # 4. Render using LibreOffice Headless
    # LibreOffice expects an output directory and generates a file with the same base name.
    output_dir = pdf_output_path.parent
    cmd = [
        soffice_path,
        "--headless",
        "--convert-to", "pdf:writer_pdf_Export",
        "--outdir", str(output_dir),
        str(render_src)
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        # Find the generated file (soffice uses input filename with .pdf)
        generated_pdf = output_dir / f"{render_src.stem}.pdf"
        if generated_pdf.exists():
            if generated_pdf != pdf_output_path:
                if pdf_output_path.exists(): pdf_output_path.unlink()
                shutil.move(str(generated_pdf), str(pdf_output_path))
        else:
            raise RuntimeError(f"LibreOffice succeeded but output file not found: {generated_pdf}")

    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice Error: {e.stderr}")
        raise RuntimeError("PDF rendering failed in the headless engine.")
    finally:
        # Drop the temp sanitized copy (never the caller's original).
        if render_src != docx_path:
            render_src.unlink(missing_ok=True)

    # 5. Validation & Mismatch Logging
    pdf_fonts = extract_pdf_fonts(pdf_output_path)
    logger.info(f"PDF embedded fonts: {pdf_fonts}")
    
    mismatches = []
    for requested in docx_fonts:
        # Simple fuzzy match: 'Arial' in 'Arial-Bold'
        if not any(requested.lower() in pf.lower() or pf.lower() in requested.lower() for pf in pdf_fonts):
            # Ignore standard MS theme fonts
            if not any(x in requested.lower() for x in ["minor", "major", "theme"]):
                mismatches.append(requested)
    
    if mismatches:
        logger.warning(f"FONT MISMATCH DETECTED: Requested {mismatches} but they were not found in the output PDF.")
        report["warnings"] = f"Font mismatch detected for: {mismatches}"
        
    return report



# LibreOffice-compatible input formats and their soffice filter names
LIBREOFFICE_INPUT_FORMATS = {
    'doc', 'docx', 'odt', 'rtf', 'html', 'htm', 'txt',
    'xls', 'xlsx', 'csv', 'ods',
    'ppt', 'pptx', 'odp', 'pps', 'ppsx',
    'docm', 'dotx', 'dot',
}

# Maps target extension to LibreOffice --convert-to argument
LIBREOFFICE_OUTPUT_MAP = {
    'pdf':  'pdf:writer_pdf_Export',
    'docx': 'docx:MS Word 2007 XML',
    'doc':  'doc:MS Word 97',
    'odt':  'odt',
    'rtf':  'rtf',
    'html': 'html:XHTML Writer File:UTF-8',
    'txt':  'txt:Text (encoded):UTF-8',
    'xlsx': 'xlsx:Calc MS Excel 2007 XML',
    'xls':  'xls:MS Excel 97',
    'csv':  'csv:Text - txt - csv (StarCalc)',
    'ods':  'ods',
    'pptx': 'pptx:Impress MS PowerPoint 2007 XML',
    'ppt':  'ppt:MS PowerPoint 97',
    'odp':  'odp',
}

def convert_document_with_libreoffice(input_path: Path, output_path: Path, target_format: str) -> dict:
    """
    Convert any LibreOffice-compatible document to target_format using soffice headless.
    Returns report dict. Raises RuntimeError on failure.
    """
    report = {"status": "success", "engine": "LibreOffice-Headless", "target_format": target_format}

    soffice_path = "soffice"
    if shutil.which("soffice") is None:
        mac_soffice = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        if mac_soffice.exists():
            soffice_path = str(mac_soffice)
        else:
            raise RuntimeError(
                "LibreOffice (soffice) not found. Install from https://www.libreoffice.org/ "
                "or run: brew install --cask libreoffice"
            )

    fmt_key = target_format.lower().lstrip('.')
    lo_filter = LIBREOFFICE_OUTPUT_MAP.get(fmt_key)
    if not lo_filter:
        raise ValueError(f"Unsupported target format: {target_format}. Supported: {list(LIBREOFFICE_OUTPUT_MAP.keys())}")

    output_dir = output_path.parent
    cmd = [
        soffice_path, "--headless",
        "--convert-to", lo_filter,
        "--outdir", str(output_dir),
        str(input_path)
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, check=True)
        logger.info(f"soffice stdout: {result.stdout.strip()}")
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice error: {e.stderr}")
        raise RuntimeError(f"LibreOffice conversion failed: {e.stderr[:200]}")
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice conversion timed out (>120s)")

    # soffice writes <stem>.<ext> in output_dir
    generated = output_dir / f"{input_path.stem}.{fmt_key}"
    if not generated.exists():
        # Try alternate: soffice may use a slightly different output name
        candidates = list(output_dir.glob(f"{input_path.stem}.*"))
        if candidates:
            generated = candidates[0]
        else:
            raise RuntimeError(f"soffice succeeded but output not found. Expected: {generated}")

    if generated != output_path:
        if output_path.exists():
            output_path.unlink()
        shutil.move(str(generated), str(output_path))

    report["output_path"] = str(output_path)
    return report


def process_txt_to_pdf(txt_path: Path, pdf_output_path: Path):
    """Convert a plain text file to PDF with full Indic shaping support."""
    report = {"status": "success"}
    
    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 if utf-8 fails
        with open(txt_path, "r", encoding="latin-1") as f:
            content = f.read()

    pdf = IndicPDF(unit="pt", format="A4")
    pdf.set_auto_page_break(auto=True, margin=36)
    pdf.add_page()
    
    # Default styling for TXT
    f_size = 12
    registered_fonts = {}
    
    # Split into paragraphs
    paragraphs = content.split('\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            pdf.ln(f_size * 1.2)
            continue
            
        # 1. Junk Stripping & Normalization
        processed_text = _clean_render_text(para_text)

        # 1.5 Script detection (all supported Indic scripts, danda-safe)
        script, registry_script = _detect_script(processed_text)

        # 2. Font Resolution: coverage-verified script font first, then fallback.
        first_cp = ord(processed_text.strip()[0]) if processed_text.strip() else None
        res_path = None
        if registry_script:
            _, res_path = _preferred_render_font(processed_text)
        if not res_path:
            res_path = font_registry.resolve_font("Normal", first_cp, script=registry_script)

        if res_path:
            f_id = res_path.stem
            if f_id not in registered_fonts:
                try:
                    pdf.add_font(f_id, "", str(res_path))
                    registered_fonts[f_id] = f_id
                except: pass
            
            pdf.set_font(f_id, size=f_size)
            
            # 3. Shaping
            pdf.set_text_shaping(use_shaping_engine=True if script else False, script=script)
            pdf.write(h=f_size * 1.3, text=processed_text)
        else:
            pdf.set_font("helvetica", size=f_size)
            pdf.write(h=f_size * 1.2, text=processed_text)
            
        pdf.ln(f_size * 1.2)

    pdf.output(str(pdf_output_path))
    return report
