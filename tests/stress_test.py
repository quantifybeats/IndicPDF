import os
import time
import requests
import sys
from pathlib import Path
from docx import Document

# Configuration
API_URL = os.environ.get("FASTAPI_BASE_URL", "http://localhost:8000")
API_KEY = os.environ.get("INDICPDF_API_KEY", "dev-key-placeholder")
HEADERS = {"X-API-Key": API_KEY}

# Paths
BASE_DIR = Path(__file__).resolve().parent
STRESS_DIR = BASE_DIR / "input" / "stress_test"
OUTPUT_DIR = BASE_DIR / "output"

def generate_heavy_doc(target_size_mb: float, filename: str):
    """Generates a DOCX file of approximately target_size_mb."""
    STRESS_DIR.mkdir(parents=True, exist_ok=True)
    file_path = STRESS_DIR / filename
    
    print(f"🛠️ Generating heavy document: {filename} (~{target_size_mb}MB)...")
    doc = Document()
    doc.add_heading(f"Stress Test: {filename}", 0)
    
    # Base text block (~1KB)
    base_text = ("IndicPDF Stress Test. " * 50) + "\n"
    
    # Approximate content needed. DOCX is compressed, so we need a LOT of text.
    # 1MB of text usually compresses down to ~50KB in DOCX.
    # To get a 24MB DOCX, we might need ~400MB of raw text.
    # That might blow up this test script's memory.
    # Let's use a loop to add paragraphs.
    
    start_time = time.time()
    for i in range(15000): # High iteration count
        doc.add_paragraph(f"Block {i}: " + base_text)
        if i % 1000 == 0:
            print(f"   Progress: {i}/15000 paragraphs added...")
            
    doc.save(file_path)
    size_mb = os.path.getsize(file_path) / (1024 * 1024)
    print(f"✅ Generated {filename}. Actual Size: {size_mb:.2f} MB. Time: {time.time() - start_time:.2f}s")
    return file_path

def run_limit_test():
    """Pushes the engine to its configured limits."""
    print("\n🔥 STARTING ENGINE LIMIT TEST 🔥")
    print("="*50)
    
    # 1. Test Single File Size Limit (Target 24.5MB)
    # We will adjust the generation to hit exactly below 25MB if possible.
    heavy_file = generate_heavy_doc(24.5, "max_size_test.docx")
    
    print(f"\n🚀 [Test 1] Uploading Single Max-Size File: {heavy_file.name}")
    start = time.time()
    with open(heavy_file, "rb") as f:
        res = requests.post(f"{API_URL}/upload", files=[("files", f)], headers=HEADERS)
    
    if res.status_code != 200:
        print(f"❌ Rejected/Failed: {res.status_code} - {res.text}")
    else:
        job_id = res.json()["jobs"][0]["job_id"]
        print(f"✅ Enqueued. Job ID: {job_id}. Polling...")
        
        while True:
            status = requests.get(f"{API_URL}/status/{job_id}", headers=HEADERS).json()
            if status["status"] == "finished":
                print(f"🎊 SUCCESS! Processed in {time.time() - start:.2f}s")
                break
            if status["status"] == "failed":
                print(f"💀 FAILED: {status.get('exc_info')}")
                break
            time.sleep(5)

    # 2. Test Batch Count Limit (10 files)
    print(f"\n🚀 [Test 2] Uploading Max Batch Size (10 files)...")
    batch_files = []
    handlers = []
    # Create 10 smaller but still significant files (~1MB each)
    for i in range(10):
        p = generate_heavy_doc(1, f"batch_stress_{i}.docx")
        f = open(p, "rb")
        handlers.append(f)
        batch_files.append(("files", (p.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")))
    
    start = time.time()
    res = requests.post(f"{API_URL}/batch/upload", files=batch_files, headers=HEADERS)
    for h in handlers: h.close()
    
    if res.status_code != 200:
        print(f"❌ Batch Rejected: {res.status_code} - {res.text}")
    else:
        batch_id = res.json()["batch_id"]
        print(f"✅ Batch Enqueued: {batch_id}. Polling...")
        while True:
            status = requests.get(f"{API_URL}/batch/status/{batch_id}", headers=HEADERS).json()
            print(f"   Progress: {status['completed']}/{status['total']} | Status: {status['status']}")
            if status["status"] == "finished":
                print(f"🎊 BATCH SUCCESS! Total Time: {time.time() - start:.2f}s")
                break
            if "failed" in status["status"]:
                print(f"💀 BATCH FAILED: {status}")
                break
            time.sleep(10)

if __name__ == "__main__":
    try:
        run_limit_test()
    except KeyboardInterrupt:
        print("\n🛑 Stress test aborted by user.")
