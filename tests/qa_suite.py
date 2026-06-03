import os
import time
import shutil
import requests
from pathlib import Path
from docx import Document

# Configuration
# Note: Ensure these environment variables are set before running
API_URL = os.environ.get("FASTAPI_BASE_URL", "http://localhost:8000")
API_KEY = os.environ.get("INDICPDF_API_KEY", "dev-key-placeholder")
HEADERS = {"X-API-Key": API_KEY}

# Paths
BASE_DIR = Path(__file__).resolve().parent
INPUT_FAST = BASE_DIR / "input" / "fast_queue"
INPUT_SLOW = BASE_DIR / "input" / "slow_queue"
OUTPUT_DIR = BASE_DIR / "output"
REPORTS_DIR = BASE_DIR / "reports"
FONTS_TELUGU_DIR = BASE_DIR / "fonts" / "telugu"
FONTS_HINDI_DIR = BASE_DIR / "fonts" / "hindi"

def generate_test_dataset():
    """Generates complex DOCX files for testing pipeline limits."""
    print("📝 Generating test dataset...")
    
    # 1. Standard Unicode Hindi (Testing ligatures and matras)
    doc1 = Document()
    doc1.add_heading("Standard Unicode Hindi Test", 0)
    run = doc1.add_paragraph().add_run("हिंदी भाषा: प्रकृति की सुंदरता। (Ligatures: प्र, वि, न्द)")
    run.font.name = "Mangal"
    doc1.save(INPUT_FAST / "01_hindi_unicode.docx")
    
    # 2. Legacy Hindi (Kruti Dev mapping failure cases - expected fail)
    doc2 = Document()
    doc2.add_heading("Legacy Hindi (Kruti Dev) Test", 0)
    run = doc2.add_paragraph().add_run("vH;kl ds fy, okD; (This is Kruti Dev mapped text)")
    run.font.name = "Kruti Dev 010"
    doc2.save(INPUT_FAST / "02_hindi_legacy_krutidev.docx")

    # 3. Mixed Language & Fonts (Telugu + Hindi + English)
    doc3 = Document()
    doc3.add_heading("Mixed Language Test", 0)
    p = doc3.add_paragraph()
    p.add_run("Telugu: నమస్కారం. ").font.name = "Gautami"
    p.add_run("Hindi: नमस्ते. ").font.name = "Mangal"
    p.add_run("English: Hello.").font.name = "Arial"
    doc3.save(INPUT_FAST / "03_mixed_languages.docx")

    # 4. Massive File (>5MB) for testing the Slow Queue routing
    print("⏳ Generating large file for slow queue testing (>5MB)...")
    doc4 = Document()
    doc4.add_heading("Slow Queue Performance Test", 0)
    base_text = "అంకితభావం మరియు సహనం (Dedication and patience) " * 100
    # Add enough content to ensure file size exceeds 5MB threshold
    for _ in range(3000): 
        doc4.add_paragraph(base_text)
    doc4.save(INPUT_SLOW / "04_large_telugu_padding.docx")
    
    # 5. Corrupt File Simulation (Magic Byte Failure)
    corrupt_path = INPUT_FAST / "05_corrupt_file.docx"
    with open(corrupt_path, "wb") as f:
        f.write(b"This is not a real DOCX file. It should fail magic byte validation because it lacks PK signature.")

def run_pipeline_test(file_path: Path):
    """Uploads file, polls status, downloads result, and returns report."""
    print(f"🚀 Testing: {file_path.name}")
    
    # 1. Upload
    try:
        with open(file_path, "rb") as f:
            # We explicitly set MIME type to ensure we hit the magic byte check on the server
            files = [("files", (file_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
            res = requests.post(f"{API_URL}/upload", files=files, headers=HEADERS)
        
        if res.status_code != 200:
            return {
                "file": file_path.name, 
                "status": "FAIL (Expected if corrupt)", 
                "error": f"HTTP {res.status_code}: {res.text[:200]}"
            }
            
        job_id = res.json()["jobs"][0]["job_id"]
    except Exception as e:
        return {"file": file_path.name, "status": "FAIL", "error": f"Upload Exception: {str(e)}"}

    # 2. Poll Status
    start_time = time.time()
    timeout = 600 # 10 minutes for slow queue
    while time.time() - start_time < timeout:
        try:
            status_res = requests.get(f"{API_URL}/status/{job_id}", headers=HEADERS)
            data = status_res.json()
            if data["status"] == "finished":
                break
            if data["status"] == "failed":
                return {"file": file_path.name, "status": "FAIL", "error": f"Processing error: {data.get('exc_info')}"}
        except Exception as e:
            return {"file": file_path.name, "status": "FAIL", "error": f"Polling Exception: {str(e)}"}
        time.sleep(2)
    else:
        return {"file": file_path.name, "status": "FAIL", "error": "Polling Timeout (10 mins)"}

    # 3. Download Result
    try:
        dl_res = requests.get(f"{API_URL}/download/{job_id}", headers=HEADERS)
        if dl_res.status_code == 200:
            out_path = OUTPUT_DIR / f"out_{file_path.stem}.pdf"
            with open(out_path, "wb") as f:
                f.write(dl_res.content)
            return {"file": file_path.name, "status": "PASS", "error": None}
        else:
            return {"file": file_path.name, "status": "FAIL", "error": f"Download failed HTTP {dl_res.status_code}"}
    except Exception as e:
        return {"file": file_path.name, "status": "FAIL", "error": f"Download Exception: {str(e)}"}

if __name__ == "__main__":
    # Note: Environment directories must exist (created by implementation step)
    if not INPUT_FAST.exists():
        print("❌ Error: Test directories missing. Run scaffold commands first.")
        exit(1)
        
    generate_test_dataset()
    
    print("\nStarting Test Execution...")
    results = []
    # Test fast queue
    for file_path in sorted(INPUT_FAST.glob("*")):
        results.append(run_pipeline_test(file_path))
    
    # Test slow queue
    for file_path in sorted(INPUT_SLOW.glob("*")):
        results.append(run_pipeline_test(file_path))
            
    print("\n" + "="*90)
    print("📊 INDICPDF QA TEST SUMMARY")
    print("="*90)
    print(f"{'FILE':<35} | {'STATUS':<25} | {'ERROR / NOTES'}")
    print("-" * 90)
    for r in results:
        icon = "✅" if "PASS" in r['status'] or "Expected" in r['status'] else "❌"
        print(f"{icon} {r['file']:<33} | {r['status']:<25} | {r['error']}")
    print("="*90)
