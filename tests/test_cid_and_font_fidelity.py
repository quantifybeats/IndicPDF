"""Regression: Indic conversion must not leak PDF text-extraction artifacts
(`(cid:N)`, `<cid:N>`, U+FFFD boxes) into rendered output, and Indic runs must
be rendered with a script-appropriate font instead of a substituted Latin one.

Root cause (proven): the DOCX->PDF render paths never stripped extraction junk
before rendering, and font substitution turned Telugu into tofu. The previous
attempt also added a toxic `resolve_cid` heuristic (chr(cid) for any CID >= 0x20)
that mapped glyph-ids to arbitrary code points -> garbage. These tests pin the
correct behaviour and lock that heuristic out for good.
"""
import sys
import unicodedata
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

import processor  # noqa: E402
from encoding_manager import encoding_manager  # noqa: E402
from font_manager import font_registry, initialize_font_registry  # noqa: E402

initialize_font_registry()

# అం(cid:9)దమై(cid:12)నది  ->  అందమైనది  ("beautiful"): the cid markers are
# spurious junk sitting between correct code points; stripping recovers Telugu.
TELUGU_WITH_CID = "అం(cid:9)దమై(cid:12)నది"
TELUGU_CLEAN = "అందమైనది"


def test_strip_handles_bracket_variants_and_replacement_char():
    messy = "అం(cid:9)ద<cid:12>మై[cid:7]న�ది"
    out = encoding_manager.strip_all_junk(messy)
    assert "cid" not in out
    assert "�" not in out
    assert "<" not in out and ">" not in out and "[" not in out and "]" not in out
    assert out == TELUGU_CLEAN


def test_clean_render_text_recovers_telugu():
    assert processor._clean_render_text(TELUGU_WITH_CID) == TELUGU_CLEAN
    # mixed bilingual line: cid gone, English preserved
    line = TELUGU_WITH_CID + ". (Telugu language is very beautiful.)"
    cleaned = processor._clean_render_text(line)
    assert "cid" not in cleaned
    assert "(Telugu language is very beautiful.)" in cleaned


def test_resolve_cid_never_invents_characters():
    """Lock out the reverted heuristic: a CID is a glyph index, NOT a code point.
    Known junk CIDs collapse to ''/' ', unknown CIDs must stay unresolved (None),
    never chr(cid)."""
    assert encoding_manager.resolve_cid(9, "AnyFont") == ""
    assert encoding_manager.resolve_cid(12, "AnyFont") == ""
    assert encoding_manager.resolve_cid(100, "AnyFont") is None  # NOT 'd'
    assert encoding_manager.resolve_cid(65, "AnyFont") is None   # NOT 'A'


def test_indic_run_gets_script_font_for_render():
    """A run whose TEXT is Telugu must resolve to an installed Telugu family,
    regardless of the (possibly Latin) font name carried by the run."""
    reg_script, path = processor._preferred_render_font(TELUGU_CLEAN)
    assert reg_script == "telugu"
    assert path is not None and path.exists()


def test_sanitize_docx_strips_cid_and_forces_telugu_font(tmp_path):
    src = tmp_path / "in.docx"
    doc = Document()
    run = doc.add_paragraph().add_run(TELUGU_WITH_CID + " (beautiful)")
    run.font.name = "Calibri"  # a Latin font name on Telugu text -> tofu risk
    doc.save(src)

    out = processor._sanitize_docx_for_render(src)
    sdoc = Document(out)
    runs = [r for p in sdoc.paragraphs for r in p.runs]
    joined = "".join(r.text for r in runs)
    assert "cid" not in joined
    assert TELUGU_CLEAN in joined
    # the Telugu-bearing run must now request an installed Telugu font
    telugu_runs = [r for r in runs if any(0x0C00 <= ord(c) <= 0x0C7F for c in r.text)]
    assert telugu_runs, "expected a Telugu run"
    for r in telugu_runs:
        assert r.font.name and font_registry.get_font_metadata(r.font.name) is not None


def test_detect_script_ignores_shared_danda():
    """Bengali/Odia/Punjabi use the Devanagari danda (U+0964). Detection must
    key off the dominant letters, not that shared punctuation mark."""
    assert processor._detect_script("বাংলা ভাষা সুন্দর।")[1] == "bengali"
    assert processor._detect_script("ଓଡ଼ିଆ ଭାଷା ସୁନ୍ଦର।")[1] == "odia"
    assert processor._detect_script("हिंदी भाषा सुंदर है।")[1] == "devanagari"
    assert processor._detect_script("தமிழ் மொழி அழகு.")[1] == "tamil"
    assert processor._detect_script("ಕನ್ನಡ ಭಾಷೆ ಸುಂದರ.")[1] == "kannada"


@pytest.mark.parametrize("script_text,reg,probe", [
    ("हिंदी भाषा", "devanagari", 0x0939),
    ("తెలుగు భాష", "telugu", 0x0C24),
    ("தமிழ் மொழி", "tamil", 0x0BAE),
    ("বাংলা ভাষা", "bengali", 0x09AC),
    ("ગુજરાતી ભાષા", "gujarati", 0x0A97),
    ("ಕನ್ನಡ ಭಾಷೆ", "kannada", 0x0C95),
    ("മലയാളം ഭാഷ", "malayalam", 0x0D2E),
    ("ଓଡ଼ିଆ ଭାଷା", "odia", 0x0B13),
])
def test_every_advertised_script_renders_with_a_covering_font(script_text, reg, probe):
    """The site advertises perfect rendering for these scripts; each must select
    a font whose cmap ACTUALLY contains the script's glyphs — broken ~21KB
    placeholder fonts (named but glyphless) must be skipped."""
    detected_reg = processor._detect_script(script_text)[1]
    assert detected_reg == reg
    rscript, path = processor._preferred_render_font(script_text)
    assert path is not None and path.exists(), f"{reg}: no font selected"
    assert processor._font_covers(path, probe), f"{reg}: {path.name} lacks the script's glyphs"


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
