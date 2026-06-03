import os
import time
import requests
from pathlib import Path
from docx import Document

# Configuration
API_URL = os.environ.get("FASTAPI_BASE_URL", "http://localhost:8000")
API_KEY = os.environ.get("INDICPDF_API_KEY", "dev-key-placeholder")
HEADERS = {"X-API-Key": API_KEY}

# Paths
BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = BASE_DIR / "input" / "batch_test"
OUTPUT_DIR = BASE_DIR / "output"

def setup_batch_test():
    """Create a set of files for batch testing."""
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    print("📝 Generating documents for batch test...")
    for i in range(1, 4):
        doc = Document()
        doc.add_heading(f"Batch Test Document {i}", 0)
        doc.add_paragraph(f"This is document number {i} in the batch.")
        doc.add_paragraph("Script: తెలుగు / हिंदी")
        file_path = INPUT_DIR / f"batch_doc_{i}.docx"
        doc.save(file_path)
        print(f"Created {file_path.name}")

def run_batch_test():
    """Uploads multiple files as a batch and verifies the merged result."""
    print("\n🚀 Starting Batch Upload Test...")
    files = []
    file_handlers = []
    
    try:
        for p in INPUT_DIR.glob("*.docx"):
            f = open(p, "rb")
            file_handlers.append(f)
            files.append(("files", (p.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")))

        # 1. Batch Upload
        res = requests.post(f"{API_URL}/batch/upload", files=files, headers=HEADERS)
        if res.status_code != 200:
            print(f"❌ Batch upload failed: {res.text}")
            return
        
        batch_data = res.json()
        batch_id = batch_data["batch_id"]
        print(f"✅ Batch enqueued! ID: {batch_id}")

        # 2. Poll Batch Status
        print(f"⏳ Polling batch status for {batch_id}...")
        start_time = time.time()
        while time.time() - start_time < 300: # 5 min timeout
            status_res = requests.get(f"{API_URL}/batch/status/{batch_id}", headers=HEADERS)
            data = status_res.json()
            
            print(f"   Status: {data['status']} | Progress: {data['completed']}/{data['total']}")
            
            if data["status"] == "finished":
                print("✅ Batch processing and merge complete!")
                break
            if "failed" in data["status"]:
                print(f"❌ Batch failed: {data}")
                return
                
            time.sleep(3)
        else:
            print("❌ Timeout waiting for batch completion.")
            return

        # 3. Download Merged Result
        print(f"📥 Downloading merged result...")
        dl_res = requests.get(f"{API_URL}/batch/download/{batch_id}", headers=HEADERS)
        if dl_res.status_code == 200:
            final_path = OUTPUT_DIR / f"merged_batch_{batch_id}.pdf"
            with open(final_path, "wb") as f:
                f.write(dl_res.content)
            print(f"✅ Success! Merged file saved to: {final_path.name}")
            print(f"   Size: {len(dl_res.content)} bytes")
        else:
            print(f"❌ Download failed: {dl_res.status_code} {dl_res.text}")

    finally:
        for f in file_handlers:
            f.close()

if __name__ == "__main__":
    setup_batch_test()
    run_batch_test()
