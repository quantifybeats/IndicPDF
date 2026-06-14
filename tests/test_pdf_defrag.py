"""Regression: PDF->DOCX must de-fragment glyph-positioned Indic PDFs.

A PDF that paints Telugu one positioned akshara at a time makes pdfminer insert a
space between every cluster and split each into its own box, so the naive
extractor emits one akshara per paragraph (`స త్ యా` ... 236 paragraphs). The
re-flow path rebuilds logical paragraphs from raw glyph coordinates, preserving
every character while restoring word spacing and paragraph structure (no OCR).
"""
import sys
import unicodedata
from pathlib import Path

import pytest
from docx import Document
from fpdf import FPDF

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

from font_manager import font_registry, initialize_font_registry  # noqa: E402
import pdf_processor as pp  # noqa: E402

initialize_font_registry()

WORDS = ["సత్యాన్వేషణ", "మరియు", "ఆత్మవిచారణ", "దైనందిన",
         "జీవితంలో", "ఎలా", "చేరుకోవాలి", "అనేది", "ముఖ్యమైన", "ప్రశ్న"]


def _clusters(s):
    out = []
    for ch in s:
        if out and (unicodedata.combining(ch) or unicodedata.category(ch).startswith("M") or ch == "్"):
            out[-1] += ch
        else:
            out.append(ch)
    return out


def _build_fragmented_pdf(path: Path):
    """A Telugu PDF where every akshara is placed individually with wide intra-
    word gaps, larger word gaps, and baseline jitter — the glyph-positioned
    pathology that shatters under pdfminer."""
    fp = str(font_registry.get_font_metadata("Noto Sans Telugu").path)
    pdf = FPDF(unit="pt", format="A4")
    pdf.add_page()
    pdf.add_font("nt", "", fp)
    pdf.set_font("nt", size=16)
    pdf.set_text_shaping(False)
    x, y, size, i = 40.0, 80.0, 16, 0
    for w in WORDS:
        for cl in _clusters(w):
            pdf.text(x, y + ((i % 3) - 1) * 1.4, cl)
            x += pdf.get_string_width(cl) + size * 0.45
            i += 1
            if x > 360:
                x, y = 40.0, y + 28
        x += size * 1.3
    pdf.output(str(path))


def test_glyph_positioned_telugu_is_defragmented(tmp_path):
    src = tmp_path / "frag.pdf"
    out = tmp_path / "frag.docx"
    _build_fragmented_pdf(src)

    report = pp.process_pdf_to_docx(src, out)
    assert report["reflowed_pages"] >= 1, "fragmentation was not detected"

    doc = Document(out)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    # Was ~one-akshara-per-paragraph; must collapse to a handful of paragraphs.
    assert len(paras) <= 4, f"still fragmented: {len(paras)} paragraphs"

    joined = " ".join(p.text for p in paras)
    tokens = joined.split()
    tiny = sum(1 for t in tokens if len(t) <= 2)
    # Most tokens must be real multi-akshara words, not single clusters.
    assert tiny / max(1, len(tokens)) < 0.4, f"too many tiny tokens: {tokens}"

    # No character loss: every Telugu code point from the source survives.
    src_chars = {c for w in WORDS for c in w if 0x0C00 <= ord(c) <= 0x0C7F}
    out_chars = {c for c in joined if 0x0C00 <= ord(c) <= 0x0C7F}
    assert src_chars <= out_chars, f"lost characters: {src_chars - out_chars}"

    # Indic runs are named with an installed Telugu font.
    for p in paras:
        for r in p.runs:
            if any(0x0C00 <= ord(c) <= 0x0C7F for c in r.text):
                assert r.font.name and font_registry.get_font_metadata(r.font.name)


def test_clean_latin_pdf_is_not_reflowed(tmp_path):
    """A normal Latin PDF must keep the fast line path (no false fragmentation)."""
    src = tmp_path / "latin.pdf"
    out = tmp_path / "latin.docx"
    pdf = FPDF(unit="pt", format="A4")
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.set_xy(40, 80)
    pdf.multi_cell(500, 16, "This is an ordinary English paragraph used to confirm "
                            "that non-Indic documents are not run through the Indic "
                            "de-fragmentation path and extract normally.")
    pdf.output(str(src))

    report = pp.process_pdf_to_docx(src, out)
    assert report["reflowed_pages"] == 0
    assert "English paragraph" in " ".join(p.text for p in Document(out).paragraphs)


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
