"""Regression: the FPDF DOCX->PDF fallback must reject zip-amplification bombs
and cap paragraph count instead of pinning CPU (QA F12)."""
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

import processor  # noqa: E402
from font_manager import initialize_font_registry  # noqa: E402

initialize_font_registry()


def _make_bomb(path: Path, repeat: int):
    """A valid .docx whose document.xml decompresses huge but zips tiny."""
    base = Document()
    base.save(path)
    body = ("<w:p><w:r><w:t>" + ("A" * 1000) + "</w:t></w:r></w:p>") * repeat
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}</w:body></w:document>"
    )
    # rewrite document.xml inside the zip
    import shutil, tempfile
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            if item == "word/document.xml":
                zout.writestr(item, xml)
            else:
                zout.writestr(item, zin.read(item))
    shutil.move(tmp, path)


def test_decompression_bomb_is_rejected(tmp_path):
    bomb = tmp_path / "bomb.docx"
    _make_bomb(bomb, repeat=25000)  # ~25 MB decompressed, tiny compressed
    decompressed = zipfile.ZipFile(bomb).getinfo("word/document.xml").file_size
    assert decompressed > processor.MAX_DOCX_XML_BYTES  # sanity: it really is over the cap
    with pytest.raises(ValueError, match="too large to render"):
        processor.process_docx_to_pdf_fpdf_fallback(bomb, tmp_path / "out.pdf")


def test_paragraph_count_is_capped(tmp_path):
    doc = Document()
    for i in range(processor.MAX_RENDER_PARAGRAPHS + 200):
        doc.add_paragraph(f"line {i}")
    src = tmp_path / "many.docx"
    doc.save(src)
    out = tmp_path / "out.pdf"
    report = processor.process_docx_to_pdf_fpdf_fallback(src, out)
    assert out.exists() and out.stat().st_size > 0
    assert report.get("truncated") is True
    assert any("truncated" in w for w in report["warnings"])
