"""Regression: Indic font resolution must pick a glyph-covering font, and the
FPDF fallback must render Indic DOCX without crashing (QA F3, F4)."""
import sys
from pathlib import Path

import pytest
from fontTools.ttLib import TTFont
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

from font_manager import initialize_font_registry, font_registry  # noqa: E402
import processor  # noqa: E402

initialize_font_registry()

# (script, a core codepoint that font MUST contain)
CORE = [("devanagari", 0x0928), ("telugu", 0x0C28), ("tamil", 0x0BA8)]


@pytest.mark.parametrize("script,cp", CORE)
def test_script_fallback_fonts_cover_their_script(script, cp):
    """Every font in a script's fallback bucket must actually render that script."""
    fonts = font_registry.script_fallback.get(script, [])
    assert fonts, f"no fallback fonts for {script}"
    not_covering = [m.path.name for m in fonts if cp not in TTFont(str(m.path)).getBestCmap()]
    assert not not_covering, f"{script} fallback has non-covering fonts: {not_covering[:5]}"


@pytest.mark.parametrize("script,cp", CORE)
def test_resolve_returns_covering_font(script, cp):
    p = font_registry.resolve_font("Normal", cp, script=script)
    assert p is not None
    assert cp in TTFont(str(p)).getBestCmap()


def test_fpdf_fallback_renders_indic_without_crash(tmp_path):
    """process_docx_to_pdf_final (FPDF fallback when LibreOffice absent) must
    not raise FPDFUnicodeEncodingException on Indic text."""
    docx = tmp_path / "hi.docx"
    d = Document()
    d.add_paragraph("नमस्ते दुनिया")   # Devanagari, run font = default ("Normal")
    d.add_paragraph("నమస్కారం")        # Telugu
    d.save(docx)
    out = tmp_path / "hi.pdf"
    report = processor.process_docx_to_pdf_final(docx, out)
    assert out.exists() and out.stat().st_size > 0
    assert not report.get("unrendered_runs"), report.get("unrendered_runs")
