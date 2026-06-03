import os
import time
import requests
import sys
from pathlib import Path
from docx import Document

# Configuration
API_URL = os.environ.get("FASTAPI_BASE_URL", "http://localhost:8000")
API_KEY = os.environ.get("INDICPDF_API_KEY", "dev-key-placeholder")
HEADERS = {"X-API-Key": API_KEY}

# Paths
BASE_DIR = Path(__file__).resolve().parent
STRESS_DIR = BASE_DIR / "input" / "memory_stress"
STRESS_DIR.mkdir(parents=True, exist_ok=True)

def generate_dense_doc(target_size_mb: float, filename: str):
    """Generates a DOCX file that is physically large (bytes) using random data."""
    file_path = STRESS_DIR / filename
    print(f"🛠️ Generating dense document: {filename} ({target_size_mb}MB)...")
    
    doc = Document()
    doc.add_heading("Memory Stress Test", 0)
    
    # Adding large chunks of text to reach size. 
    # To prevent compression, we use random-ish strings.
    import random
    import string
    
    chars = string.ascii_letters + string.digits
    
    current_size = 0
    # We target slightly less because DOCX adds overhead.
    # 25MB DOCX is very large.
    while current_size < target_size_mb * 1024 * 1024:
        # Create a block of 100KB unique text
        block = ''.join(random.choices(chars, k=100000))
        doc.add_paragraph(block)
        # Approximate current size (this is very loose)
        current_size += 100000 
        if current_size % 1000000 == 0:
            print(f"   Generated ~{current_size/1000000:.1f} MB of raw text...")

    doc.save(file_path)
    actual_size = os.path.getsize(file_path) / (1024 * 1024)
    print(f"✅ Generated {filename}. Actual Size: {actual_size:.2f} MB")
    return file_path

def run_memory_attack():
    """Simulates a concurrent upload of max-size files to trigger OOM."""
    print("\n💀 STARTING MEMORY EXHAUSTION TEST 💀")
    print("="*50)
    
    # 1. Generate 4 files of ~24MB each.
    # If worker concurrency is 2, 2x50MB = 100MB RAM spike.
    # If worker concurrency was default (CPU count), this would crash a 512MB box.
    files = []
    for i in range(4):
        files.append(generate_dense_doc(24, f"dense_{i}.docx"))
    
    print(f"\n🚀 Uploading 4 Dense Files simultaneously...")
    job_ids = []
    for p in files:
        with open(p, "rb") as f:
            res = requests.post(f"{API_URL}/upload", files=[("files", f)], headers=HEADERS)
            if res.status_code == 200:
                jid = res.json()["jobs"][0]["job_id"]
                job_ids.append(jid)
                print(f"   Enqueued: {jid}")
            else:
                print(f"   Rejected: {res.status_code} - {res.text}")

    print("\nPolling for results. If worker crashes, we will see 'failed' or infinite polling.")
    start = time.time()
    while job_ids:
        for jid in list(job_ids):
            status = requests.get(f"{API_URL}/status/{jid}", headers=HEADERS).json()
            if status["status"] == "finished":
                print(f"✅ {jid} Finished in {time.time() - start:.1f}s")
                job_ids.remove(jid)
            elif status["status"] == "failed":
                print(f"❌ {jid} Failed: {status.get('exc_info')}")
                job_ids.remove(jid)
        time.sleep(5)
    
    print("\n🎊 Memory Stress Test Complete.")

if __name__ == "__main__":
    run_memory_attack()
